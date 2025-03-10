import fitz  # PyMuPDF
import re
import json
import docx
import requests
import time
import os
from docx import Document
from typing import Dict, List
from langdetect import detect

class ProposalGenerator:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.api_url = "https://api.deepseek.com/v1/chat/completions"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        self.chunk_size = 2500
        self.request_delay = 2.0
        self.context_window = 3500
        self.language = "es"
        self.temp_file = "temp_proposal.docx"

    def log_progress(self, message: str):
        timestamp = time.strftime("%H:%M:%S", time.localtime())
        print(f"[{timestamp}] {message}")

    def extract_titles_from_pdf(self, pdf_path: str) -> List[str]:
        """Extract structured titles from the PDF document."""
        doc = fitz.open(pdf_path)
        text = "".join([page.get_text("text") + "\n" for page in doc])
        section_pattern = re.compile(r"^(?:\d+(\.\d+)*|[A-Z]+|[IVXLCDM]+)\.-+\s.+")
        titles = [line.strip() for line in text.split("\n") if section_pattern.match(line.strip())]
        self.log_progress(f"Títulos extraídos del PDF: {titles}")
        return titles

    def post_process(self, text: str) -> str:
        """Format content to remove redundancy and improve readability."""
        paragraphs = []
        seen = set()
        for para in text.split('\n'):
            key = re.sub(r'\s+', '', para).lower()[:50]
            if key and key not in seen:
                seen.add(key)
                paragraphs.append(para.replace("**", "").strip())
        return '\n\n'.join(paragraphs)

    def analyze_section(self, section: str, requirements: str, proposal_data: Dict) -> str:
        """Enhanced content generation with compliance considerations."""
        for attempt in range(3):
            try:
                response = requests.post(
                    self.api_url,
                    json={
                        "model": "deepseek-chat",
                        "messages": [{
                            "role": "user",
                            "content": f"""Genera una sección de propuesta en español para: {section}
                            Requisitos técnicos: {requirements[:2000]}
                            Componentes existentes: {proposal_data}
                            - Incluye detalles sobre certificaciones (ENS, UNE-EN 50600, ISO 27001)
                            - Asegura que la infraestructura cumple con 4 CPU, 12GB RAM, 100GB SSD, transferencia 300Mb/s
                            - Ajusta los SLA a: 99.3% disponibilidad, resolución de urgencias en 6h, estándar en 24h
                            - Asegura el cumplimiento con RGPD, Ley Orgánica 3/2018 y medidas de protección de datos
                            - Considera formación y transferencia de conocimiento al finalizar el contrato
                            - Asegura que el CMS será desarrollado a medida y será compatible con PHP 8, Laravel 5.4, MariaDB 10.3.39
                            - Evita repeticiones innecesarias mientras aseguras que todos los requisitos están cubiertos.
                            Salida SOLO en contenido de la sección, sin preámbulos ni instrucciones."""
                        }],
                        "temperature": 0.25,
                        "max_tokens": 2000
                    },
                    headers=self.headers,
                    timeout=900
                )
                self.log_progress(f"Respuesta API: {response.status_code}, {response.text[:500]}")
                if response.status_code == 200:
                    content = response.json()['choices'][0]['message']['content']
                    return self.post_process(content)
                time.sleep(2 ** attempt)
            except Exception as e:
                self.log_progress(f"Intento {attempt+1} fallido: {str(e)}")
        return ""

    def generate_proposal(self, tech_spec_path: str, pdf_path: str, output_path: str):
        """Generate proposal based on extracted tender structure."""
        try:
            self.log_progress("Iniciando la generación de la propuesta en español")
            structured_titles = self.extract_titles_from_pdf(pdf_path)
            if not structured_titles:
                self.log_progress("⚠️ No se encontraron títulos en el PDF. Verifique el formato del documento.")
                return
            tech_doc = Document(tech_spec_path)
            requirements = "\n".join(p.text.strip() for p in tech_doc.paragraphs if p.text.strip())
            
            # Load existing or create a new document
            if os.path.exists(self.temp_file):
                output_doc = Document(self.temp_file)
                self.log_progress("Se retomará la generación desde el archivo temporal.")
            else:
                output_doc = Document()

            generated_sections = {p.text.strip().lower() for p in output_doc.paragraphs}
            
            for section in structured_titles:
                if section.lower() in generated_sections:
                    self.log_progress(f"⏭️ Se omite: {section} (ya procesado)")
                    continue
                self.log_progress(f"Procesando: {section}")
                content = self.analyze_section(section, requirements, {})
                if content:
                    output_doc.add_heading(section, level=2)
                    output_doc.add_paragraph(content)
                    generated_sections.add(section.lower())
                    output_doc.save(self.temp_file)  # Save progress incrementally
                    time.sleep(self.request_delay)
            
            if len(output_doc.paragraphs) == 0:
                self.log_progress("⚠️ No se generó contenido para la propuesta.")
                return

            output_doc.save(output_path)
            os.remove(self.temp_file)  # Cleanup temp file after completion
            self.log_progress(f"Propuesta final guardada en {output_path}")
        except Exception as e:
            self.log_progress(f"Error: {str(e)}")
            raise

if __name__ == "__main__":
    API_KEY = "your_api_key_here"
    generator = ProposalGenerator(API_KEY)
    generator.generate_proposal(
        tech_spec_path="Pliego_Tecnico_Limpio.docx",
        pdf_path="PPT.PDF",
        output_path="Final_Proposal.docx"
    )


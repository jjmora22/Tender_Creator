import os
import re
import docx
import requests
import time
from docx import Document
from typing import Dict, List

class ProposalGenerator:
    def __init__(self):
        """Initialize the Proposal Generator with API settings."""
        self.api_key = os.getenv("DEEPSEEK_API_KEY")  # Get API key from environment variable
        if not self.api_key:
            raise ValueError("API key is missing. Set DEEPSEEK_API_KEY as an environment variable.")
        
        self.api_url = "https://api.deepseek.com/v1/chat/completions"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        self.request_delay = 1.5  # Delay in seconds between API calls
        self.timeout = 800  # API request timeout in seconds
        self.retries = 2  # Number of retries for failed API calls

    def log_progress(self, message: str):
        """Logs progress messages with timestamps."""
        timestamp = time.strftime("%H:%M:%S", time.localtime())
        print(f"[{timestamp}] {message}")

    def analyze_requirement(self, requirement: str, proposal_data: Dict) -> str:
        """Processes a technical requirement and generates proposal text."""
        self.log_progress("Starting API analysis for requirement...")
        
        for attempt in range(self.retries + 1):
            try:
                response = requests.post(
                    self.api_url,
                    json={
                        "model": "deepseek-chat",
                        "messages": [{
                            "role": "user",
                            "content": f"""Transform this technical requirement into proposal text using these components: {proposal_data}.
                            Requirement: {requirement}
                            Respond ONLY with the proposal text, no markdown."""
                        }],
                        "temperature": 0.2,
                        "max_tokens": 1000
                    },
                    headers=self.headers,
                    timeout=self.timeout
                )

                self.log_progress(f"API response received (Status: {response.status_code})")

                # Check for successful API response
                if response.status_code != 200:
                    self.log_progress(f"API Error: Received status {response.status_code}")
                    return f"{requirement} [API Error: Status {response.status_code}]"

                result = response.json()['choices'][0]['message']['content']
                self.log_progress("Successfully processed requirement")
                return result

            except requests.exceptions.RequestException as e:
                self.log_progress(f"Attempt {attempt+1} failed: {str(e)}")
                if attempt < self.retries:
                    time.sleep(2 ** attempt)  # Exponential backoff for retries
                else:
                    return f"{requirement} [Error: {str(e)}]"

            time.sleep(self.request_delay)
        
        return requirement

    def extract_proposal_components(self, doc_path: str) -> Dict:
        """Extracts components from a proposal document for structured transformation."""
        self.log_progress("Starting proposal component extraction...")
        doc = Document(doc_path)
        components = {
            "sections": {"General": []},  # Default section
            "raw": []
        }    
        current_section = "General"

        try:
            for para in doc.paragraphs:
                text = para.text.strip()
                components["raw"].append(text)
            
                # Detect and organize sections based on headings
                if text.startswith(('## ', '##', '#')):  # Multiple heading formats
                    current_section = text.replace('#', '').strip()
                    components["sections"][current_section] = []
                else:
                    if text:  # Skip empty paragraphs
                        components["sections"].setdefault(current_section, []).append(text)
        
            self.log_progress(f"Extracted {len(components['sections'])} sections")
            return components
        except Exception as e:
            self.log_progress(f"Extraction error: {str(e)}")
            return components  # Return partial results

    def generate_proposal(self, tech_spec_path: str, proposal_path: str, output_path: str):
        """Generates a proposal document based on technical specifications and extracted components."""
        try:
            self.log_progress("Starting proposal generation process")
            
            # Phase 1: Extract Proposal Components
            self.log_progress("Phase 1: Extracting proposal components")
            proposal_data = self.extract_proposal_components(proposal_path)
            
            if "error" in proposal_data:
                raise ValueError("Failed to extract proposal components")

            # Phase 2: Process Technical Specifications
            self.log_progress("Phase 2: Processing technical specifications")
            tech_doc = Document(tech_spec_path)
            output_doc = Document()
            
            total_paragraphs = len(tech_doc.paragraphs)
            self.log_progress(f"Found {total_paragraphs} paragraphs to process")

            # Phase 3: Transform Paragraphs
            self.log_progress("Phase 3: Transforming requirements")
            for idx, para in enumerate(tech_doc.paragraphs):
                self.log_progress(f"Processing paragraph {idx+1}/{total_paragraphs}")
                transformed = self.analyze_requirement(para.text, proposal_data)
                output_doc.add_paragraph(transformed)
                time.sleep(self.request_delay)

            # Phase 4: Save Final Document
            self.log_progress("Phase 4: Saving final document")
            output_doc.save(output_path)
            self.log_progress(f"Proposal successfully saved to {output_path}")

        except Exception as e:
            self.log_progress(f"Critical error: {str(e)}")
            raise

if __name__ == "__main__":
    try:
        generator = ProposalGenerator()
        generator.generate_proposal(
            tech_spec_path="Pliego_Tecnico_Limpio.docx",
            proposal_path="Resumen_Propuesta.docx",
            output_path="propuesta_vers1b.docx"
        )
    except Exception as e:
        print(f"\nFatal error occurred: {str(e)}")
        print("Please check:")
        print("1. API key validity")
        print("2. Internet connection")
        print("3. Document accessibility")

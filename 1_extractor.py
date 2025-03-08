import fitz  # PyMuPDF for reading PDFs
import re     # Regular expressions for pattern matching
import spacy  # Natural language processing library
from odf.opendocument import load  # For loading ODT files
from odf.text import P             # For accessing text elements in ODT files
from docx import Document          # For creating and saving Word documents

# Load the spaCy English model.
# Ensure you have installed it with:
#   pip install spacy
#   python -m spacy download en_core_web_sm
nlp = spacy.load("en_core_web_sm")

def extract_text_from_pdf(pdf_path):
    """
    Extracts text from a PDF file while ignoring metadata, headers, footers, and digital seals.
    
    Parameters:
        pdf_path (str): The file path of the PDF document.
        
    Returns:
        str: A string containing the cleaned and concatenated text from the PDF.
    """
    # Open the PDF document.
    doc = fitz.open(pdf_path)
    extracted_text = []
    
    # Iterate through each page of the PDF.
    for page in doc:
        # Extract text blocks with coordinates from the page.
        blocks = page.get_text("blocks")
        for block in blocks:
            x0, y0, x1, y1, text, _, _ = block  # Unpack block data: coordinates and text
            
            # Determine the page height.
            page_height = page.rect.height
            # Check if the text is not in the header or footer area (by margin threshold)
            if y0 > 80 and y1 < (page_height - 80):
                # Exclude text matching certain patterns (common in metadata and signatures).
                if not re.search(r'COPIA AUTÉNTICA|Verificación|firmado por|Fecha/hora|Cargo:', text, re.IGNORECASE):
                    extracted_text.append(text.strip())
    
    # Remove lines that are likely digital seals (long hexadecimal strings)
    cleaned_text = []
    for line in extracted_text:
        if not re.match(r'^[0-9a-fA-F]{20,}$', line.strip()):
            cleaned_text.append(line)
    
    return "\n".join(cleaned_text)

def extract_text_from_odt(odt_path):
    """
    Extracts text from an ODT (OpenDocument Text) file while preserving the structure.
    
    Parameters:
        odt_path (str): The file path of the ODT document.
        
    Returns:
        str: A string containing the concatenated text from the ODT.
    """
    text = ""
    # Load the ODT document.
    doc = load(odt_path)
    # Iterate through all text paragraph elements.
    for element in doc.getElementsByType(P):
        if element.childNodes:
            # Concatenate text from all text nodes.
            text += "".join(node.data for node in element.childNodes if node.nodeType == node.TEXT_NODE) + " "
    return text.strip()

def clean_and_format_text(text):
    """
    Cleans and formats the provided text using spaCy and regular expressions.
    It fixes sentence boundaries, merges split lines, and normalizes bullet points and headings.
    
    Parameters:
        text (str): The raw text to be cleaned.
        
    Returns:
        str: The cleaned and formatted text.
    """
    doc = nlp(text)
    formatted_lines = []
    prev_line = ""
    
    # Process each sentence as determined by spaCy.
    for sent in doc.sents:
        line = sent.text.strip()
        
        # Merge lines that were split incorrectly (if the previous line does not end with punctuation).
        if prev_line and not prev_line.endswith(('.', ':', '!', '?')):
            formatted_lines[-1] += f" {line}"
        else:
            formatted_lines.append(line)
        prev_line = line
        
        # Normalize bullet points: Replace various bullet symbols with a uniform dash.
        line = re.sub(r'•\s*', '- ', line)
        line = re.sub(r'▪\s*', '- ', line)
        line = re.sub(r'◦\s*', '  - ', line)  # For sub-level bullets
        
        # Normalize headings: Bold numbered sections and markdown-style headings.
        line = re.sub(r'^(\d+\.\s.+)', r'**\1**', line)
        line = re.sub(r'^(##\s*.+)', r'**\1**', line)
    
    return "\n".join(formatted_lines)

def create_formatted_word_document(text, output_path):
    """
    Creates a Word document (.docx) with the formatted text.
    
    Parameters:
        text (str): The text to be added to the document.
        output_path (str): The file path to save the generated Word document.
        
    Returns:
        str: The output path of the saved document.
    """
    doc = Document()
    
    # Iterate through each line of the formatted text.
    for line in text.split("\n"):
        # If the line is formatted as a title (enclosed in '**'), add it as a heading.
        if line.startswith("**") and line.endswith("**"):
            doc.add_heading(line.strip('**'), level=2)
        else:
            # Otherwise, add it as a normal paragraph.
            doc.add_paragraph(line)
    
    # Save the Word document to the specified output path.
    doc.save(output_path)
    return output_path

# Define input file paths.
pdf_ppt = "PPT.pdf"       # PDF file containing Technical Specifications.
pdf_pcap = "PCAP.pdf"     # PDF file containing Administrative Specifications.
odt_anexos = "ANEXOS.odt"  # ODT file containing annex documents.

# Extract raw text from the input files.
text_ppt = extract_text_from_pdf(pdf_ppt)
text_pcap = extract_text_from_pdf(pdf_pcap)
text_anexos = extract_text_from_odt(odt_anexos)

# Clean and format the text extracted from the PPT PDF.
formatted_text = clean_and_format_text(text_ppt)

# Define the output file path for the formatted Word document.
output_word = "Pliego_Tecnico_Limpio_LocalNLP.docx"

# Create and save the formatted Word document.
create_formatted_word_document(formatted_text, output_word)
